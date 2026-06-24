import os
import re
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level, space_before=12, space_after=6):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    # Customize style
    run = heading.runs[0]
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(26, 54, 93)  # Deep Blue
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(44, 82, 130)  # Slate Blue
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(74, 85, 104)  # Charcoal
        
    return heading

def parse_evaluation_report(report_path):
    print(f"Parsing evaluation report: {report_path}")
    with open(report_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find False Positives section
    fp_section = re.search(r'### False Positives \(\d+\)(.*?)(### False Negatives|\Z)', content, re.DOTALL)
    if not fp_section:
        raise ValueError("Could not find False Positives section in evaluation report.")
        
    fp_text = fp_section.group(1)
    
    # Each patient block starts with **Patient:**
    patient_blocks = re.split(r'\*\*Patient:\*\*\s*', fp_text)
    
    false_positives = []
    for block in patient_blocks:
        block = block.strip()
        if not block:
            continue
            
        # Parse grid/ID
        lines = block.split('\n')
        grid = lines[0].strip()
        
        # Parse GT/Agent
        gt_line = ""
        comment_line = ""
        reasoning_lines = []
        
        current_field = None
        for line in lines[1:]:
            line_str = line.strip()
            if line_str.startswith('**GT:**'):
                gt_line = line_str
            elif line_str.startswith('**Comment:**'):
                comment_line = line_str.replace('**Comment:**', '').strip()
            elif line_str.startswith('**Agent Reasoning:**'):
                current_field = 'reasoning'
                reasoning_lines.append(line_str.replace('**Agent Reasoning:**', '').strip())
            elif current_field == 'reasoning' and line_str:
                reasoning_lines.append(line_str)
                
        reasoning = " ".join(reasoning_lines)
        
        # Extract GT and Agent labels
        gt_match = re.search(r'\*\*GT:\*\*\s*([^\s|]+)', gt_line)
        agent_match = re.search(r'\*\*Agent:\*\*\s*([^\s|]+)', gt_line)
        orig_match = re.search(r'\*\*Original GT:\*\*\s*(.*)', gt_line)
        
        gt_label = gt_match.group(1) if gt_match else 'Negative'
        agent_label = agent_match.group(1) if agent_match else 'Positive'
        orig_gt = orig_match.group(1) if orig_match else 'No'
        
        false_positives.append({
            'grid': grid,
            'gt': gt_label,
            'agent': agent_label,
            'original_gt': orig_gt,
            'comment': comment_line,
            'reasoning': reasoning
        })
        
    print(f"Successfully parsed {len(false_positives)} false positive patients.")
    return false_positives

def parse_patient_markdown(md_path):
    if not os.path.exists(md_path):
        print(f"Warning: Markdown file {md_path} not found.")
        return {'labs': [], 'notes': []}
        
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    # Parse Labs section
    labs = []
    labs_match = re.search(r'## Labs\n(.*?)(## Medical Notes|\Z)', content, re.DOTALL)
    if labs_match:
        labs_text = labs_match.group(1).strip()
        for line in labs_text.split('\n'):
            line = line.strip()
            if line.startswith('-'):
                # Extract clean lab details
                labs.append(line.replace('- ', '').strip())
            elif line and line != 'No labs available.':
                labs.append(line)
                
    # Parse Medical Notes section
    notes = []
    notes_match = re.search(r'## Medical Notes\n(.*)', content, re.DOTALL)
    if notes_match:
        notes_text = notes_match.group(1).strip()
        # Find all note sub-headers: ### [date] title
        note_splits = re.split(r'###\s+\[', notes_text)
        for split in note_splits:
            split = split.strip()
            if not split:
                continue
            
            # Format is: date_str] title\n**Source:** source\n\nbody
            parts = split.split(']', 1)
            if len(parts) < 2:
                continue
            date_str = parts[0].strip()
            rest = parts[1].strip()
            
            lines = rest.split('\n')
            title = lines[0].strip()
            
            source = "Unknown"
            body_start = 1
            if len(lines) > 1 and lines[1].strip().startswith('**Source:**'):
                source = lines[1].replace('**Source:**', '').strip()
                body_start = 2
                
            body = "\n".join(lines[body_start:]).strip()
            
            notes.append({
                'date': date_str,
                'title': title,
                'source': source,
                'body': body
            })
            
    # Load and use the exact popular words from the agentic system
    import yaml
    yaml_path = "/home/biand/Projects/Celiac_BioVU/data/celiac_keywords_latest.yaml"
    with open(yaml_path, 'r', encoding='utf-8') as f:
        keywords_data = yaml.safe_load(f)
    popular_words = keywords_data.get('popular_words', [])
    
    filtered_notes = []
    if popular_words:
        # Match using same case-insensitive contains logic as the agentic system (contains substring)
        pattern = re.compile('|'.join([re.escape(w) for w in popular_words]), re.IGNORECASE)
        for note in notes:
            text_to_check = f"{note['title']} {note['source']} {note['body']}"
            if pattern.search(text_to_check):
                filtered_notes.append(note)
    else:
        filtered_notes = notes
        
    # Fallback to avoid empty records if none matched (should rarely happen)
    if not filtered_notes and notes:
        filtered_notes = notes
        
    return {'labs': labs, 'notes': filtered_notes}

def style_document(doc):
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 37, 41)  # Off-black/Charcoal
    
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

def add_patient_to_doc(doc, fp, clinical_data):
    # 1. Heading: Patient Grid ID
    add_styled_heading(doc, f"Clinical Review Packet: Patient {fp['grid']}", level=1, space_before=18, space_after=12)
    
    # 2. Metadata Table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Shading Accent 1'
    table.autofit = False
    
    # Set column widths
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.7)
    
    headers = [
        ("Patient Identifier", fp['grid']),
        ("Ground Truth Classification", f"{fp['gt']} (Confirmed Celiac: {fp['original_gt']})"),
        ("Model Classification", f"{fp['agent']} (False Positive Classification)"),
        ("Clinical Review Comments", fp['comment']),
        ("Algorithm Decision Logic", fp['reasoning'])
    ]
    
    for idx, (label, val) in enumerate(headers):
        row = table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        # Style label cell
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        run_lbl.bold = True
        run_lbl.font.size = Pt(10)
        run_lbl.font.color.rgb = RGBColor(26, 54, 93)
        set_cell_background(cell_lbl, "F0F4F8")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        
        # Style value cell
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        run_val.font.size = Pt(10)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # 3. Labs Section
    add_styled_heading(doc, "Laboratory Results", level=2, space_before=14, space_after=6)
    if clinical_data['labs']:
        for lab in clinical_data['labs']:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Find date in square brackets
            match = re.match(r'^\[(.*?)\]', lab)
            if match:
                date_str = match.group(1)
                rest = lab[len(date_str)+2:].strip()
                r_date = p.add_run(f"[{date_str}] ")
                r_date.bold = True
                p.add_run(rest)
            else:
                p.add_run(lab)
    else:
        p = doc.add_paragraph()
        r = p.add_run("No relevant laboratory values identified in the electronic medical record.")
        r.italic = True
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # 4. Clinical Notes Section
    add_styled_heading(doc, "Longitudinal Medical Notes", level=2, space_before=14, space_after=6)
    
    if clinical_data['notes']:
        # Sort notes chronologically (just in case)
        for idx, note in enumerate(clinical_data['notes'], 1):
            # Note Title & Metadata
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(8)
            p_title.paragraph_format.space_after = Pt(2)
            p_title.paragraph_format.keep_with_next = True
            
            r_num = p_title.add_run(f"Note {idx}: [{note['date']}] ")
            r_num.bold = True
            r_num.font.color.rgb = RGBColor(44, 82, 130)
            
            r_title = p_title.add_run(note['title'])
            r_title.bold = True
            
            p_src = doc.add_paragraph()
            p_src.paragraph_format.space_after = Pt(4)
            p_src.paragraph_format.keep_with_next = True
            r_src_lbl = p_src.add_run("Document Source: ")
            r_src_lbl.font.size = Pt(9.5)
            r_src_lbl.bold = True
            r_src_val = p_src.add_run(note['source'])
            r_src_val.font.size = Pt(9.5)
            r_src_val.italic = True
            
            # Note Text - Split by double newlines or blank lines to reduce paragraph count significantly
            body_text = note['body'].strip()
            blocks = re.split(r'\n\s*\n', body_text)
            for block in blocks:
                block_clean = block.strip()
                if not block_clean:
                    continue
                # Split extremely long blocks into manageable paragraph chunks (safety limit 15k chars)
                chunk_size = 15000
                chunks = [block_clean[i:i+chunk_size] for i in range(0, len(block_clean), chunk_size)]
                for chunk in chunks:
                    p_body = doc.add_paragraph()
                    p_body.paragraph_format.left_indent = Inches(0.25)
                    p_body.paragraph_format.space_after = Pt(4)
                    r_body = p_body.add_run(chunk)
                    r_body.font.size = Pt(9.5)
                
            # Add a subtle divider between notes
            if idx < len(clinical_data['notes']):
                p_div = doc.add_paragraph()
                p_div.paragraph_format.space_before = Pt(6)
                p_div.paragraph_format.space_after = Pt(6)
                p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_div = p_div.add_run("·  ·  ·  ·  ·  ·  ·  ·  ·  ·  ·  ·  ·  ·  ·  ·  ·")
                r_div.font.color.rgb = RGBColor(160, 174, 192)
                r_div.font.size = Pt(9)
    else:
        p = doc.add_paragraph()
        r = p.add_run("No clinical notes identified in the electronic medical record.")
        r.italic = True

def main():
    report_path = "/home/biand/Projects/Celiac_BioVU/results/celiac_agent_v3/summary/evaluation_report.md"
    dataset_dir = "/home/biand/Projects/Celiac_BioVU/data/ehr_markdown_dataset"
    out_dir = "/home/biand/Projects/Celiac_BioVU/results/celiac_agent_v3/clinical_review"
    
    os.makedirs(out_dir, exist_ok=True)
    
    try:
        false_positives = parse_evaluation_report(report_path)
    except Exception as e:
        print(f"Error parsing report: {e}")
        return
        
    # 1. Create a combined document containing all false positives
    combined_doc = docx.Document()
    style_document(combined_doc)
    
    # Title Page for the Combined Review Package
    p_title = combined_doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(120)
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("CLINICAL REVIEW PACKAGE\nCeliac Disease Diagnostic Classification")
    r_title.font.size = Pt(26)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(26, 54, 93)
    
    p_sub = combined_doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(40)
    r_sub = p_sub.add_run("Comprehensive Audit of False Positive Electronic Health Records (EHR)\nCeliac Agent Version 3 Evaluation")
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = RGBColor(74, 85, 104)
    
    p_meta = combined_doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_before = Pt(150)
    r_meta = p_meta.add_run("Prepared for Clinical Reviewers and Medical Evaluators\nTotal Records: 38 False Positives\nDate: June 2026")
    r_meta.font.size = Pt(11)
    r_meta.italic = True
    
    combined_doc.add_page_break()
    
    # 2. Iterate through each patient and generate files
    for idx, fp in enumerate(false_positives, 1):
        print(f"[{idx}/{len(false_positives)}] Processing Patient {fp['grid']}...")
        
        md_file = os.path.join(dataset_dir, f"{fp['grid']}.md")
        clinical_data = parse_patient_markdown(md_file)
        
        # Create individual document
        indiv_doc = docx.Document()
        style_document(indiv_doc)
        add_patient_to_doc(indiv_doc, fp, clinical_data)
        
        # Save individual document
        indiv_path = os.path.join(out_dir, f"Patient_{fp['grid']}.docx")
        indiv_doc.save(indiv_path)
        
        # Add to combined document using fast XML body append to bypass O(N^2) paragraph insertions
        from copy import deepcopy
        for element in indiv_doc.element.body:
            if element.tag.endswith('sectPr'):
                continue
            combined_doc.element.body.append(deepcopy(element))
        
        # Add page break in combined if not the last patient
        if idx < len(false_positives):
            combined_doc.add_page_break()
            
    # Save combined document
    combined_path = os.path.join(out_dir, "All_False_Positives_Review_Package.docx")
    combined_doc.save(combined_path)
    print(f"\nSaved combined review package to: {combined_path}")
    print(f"Generated {len(false_positives)} individual clinician-review documents in: {out_dir}")

if __name__ == '__main__':
    main()
